import os
from pathlib import Path
import tempfile
import cv2
from docx import Document
from docx.shared import Inches
from PIL import Image


def process_videos(base_folder: str):
    """
    对 base_folder 及其子目录：
      - 如果没有 .mp4 或 .mov 文件，则生成 docx 提示该生未上传汇报视频。
      - 否则对每个视频（.mp4 和 .mov）：
          1. 获取时长
          2. 按 16 等分时间点截帧并插入文档
          3. 保存 "<video_stem>_summary.docx"
          4. 删除原始视频
    """
    base_dir = Path(base_folder)
    if not base_dir.is_dir():
        raise ValueError(f"“{base_folder}” 不是有效目录")

    for folder in base_dir.rglob('*'):  # 修改为 rglob('*')，扫描所有子目录
        if not folder.is_dir():
            continue

        video_files = list(folder.glob("*.mp4")) + list(folder.glob("*.mov"))

        # 如果没有视频文件，生成提示文档
        if not video_files:
            doc = Document()
            doc.add_paragraph(f"文件夹“{folder.name}”中未发现视频。")
            doc_path = folder / "未上传视频汇报.docx"
            doc.save(str(doc_path))
            print(f"生成提示文档：{doc_path}")
            continue

        # 有视频，逐个处理
        for video_path in video_files:
            cap = cv2.VideoCapture(str(video_path))
            fps = cap.get(cv2.CAP_PROP_FPS) or 1
            total_frames = cap.get(cv2.CAP_PROP_FRAME_COUNT)
            duration = total_frames / fps
            timestamps = [(i + 1) * duration / 17 for i in range(16)]

            doc = Document()
            doc.add_paragraph(f"视频文件：{video_path.name}")
            doc.add_paragraph(f"视频时长：{duration:.2f} 秒")

            temp_files = []
            for idx, t in enumerate(timestamps):
                frame_no = int(t * fps)
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_no)
                ret, frame = cap.read()
                if not ret:
                    continue

                # 截图并保存为临时文件
                img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                img.save(tmp.name)
                tmp.close()
                temp_files.append(tmp.name)

                # 将截图插入到文档
                doc.add_picture(tmp.name, width=Inches(2))

            cap.release()

            # 保存文档
            docx_path = folder / f"{video_path.stem}_summary.docx"
            doc.save(str(docx_path))
            print(f"生成文档：{docx_path}")

            # 删除原视频和临时文件
            try:
                video_path.unlink()
                print(f"删除视频文件：{video_path}")
            except OSError as e:
                print(f"删除视频文件失败：{video_path}，错误：{e}")

            for f in temp_files:
                try:
                    os.remove(f)
                except OSError:
                    pass


def a02_11_video_summary(base_folder: str):
    """
    执行视频处理流程，包括 .mp4 和 .mov 文件格式。
    """
    process_videos(base_folder)
